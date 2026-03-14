"""Document parsing service - extract text from PDF/DOCX/MD files."""

import re
import fitz  # PyMuPDF
from docx import Document
from pathlib import Path


def parse_document(content: bytes, filename: str) -> dict:
    """Parse a document and return structured text with paragraph numbering.

    Returns:
        {
            "text": str,           # Full text with paragraph numbers
            "paragraphs": list,    # List of paragraph strings
            "headers": list,       # List of {"level": int, "text": str, "para_index": int}
            "page_count": int,
        }
    """
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return _parse_pdf(content)
    elif ext == ".docx":
        return _parse_docx(content)
    elif ext == ".md":
        return _parse_markdown(content)
    else:
        raise ValueError(f"Unsupported file format: {ext}")


def _parse_pdf(content: bytes) -> dict:
    """Extract text from PDF preserving structure via font size heuristics."""
    doc = fitz.open(stream=content, filetype="pdf")
    paragraphs = []
    headers = []

    for page in doc:
        blocks = page.get_text("dict")["blocks"]
        for block in blocks:
            if block.get("type") != 0:  # text block only
                continue
            for line_group in block.get("lines", []):
                spans = line_group.get("spans", [])
                if not spans:
                    continue
                text = "".join(s["text"] for s in spans).strip()
                if not text:
                    continue

                # Detect headers by font size (>= 14pt as header)
                max_size = max(s["size"] for s in spans)
                para_index = len(paragraphs)
                paragraphs.append(text)

                if max_size >= 14:
                    level = 1 if max_size >= 18 else 2 if max_size >= 16 else 3
                    headers.append({"level": level, "text": text, "para_index": para_index})

    page_count = doc.page_count
    doc.close()

    numbered_text = _build_numbered_text(paragraphs)
    return {
        "text": numbered_text,
        "paragraphs": paragraphs,
        "headers": headers,
        "page_count": page_count,
    }


def _parse_docx(content: bytes) -> dict:
    """Extract text from DOCX preserving heading styles."""
    import io
    doc = Document(io.BytesIO(content))
    paragraphs = []
    headers = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        para_index = len(paragraphs)
        paragraphs.append(text)

        # Detect headings by style name
        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading"):
            try:
                level = int(style_name.replace("Heading", "").strip())
            except ValueError:
                level = 1
            headers.append({"level": level, "text": text, "para_index": para_index})

    # Also extract table content
    for table in doc.tables:
        table_lines = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_lines.append(" | ".join(cells))
        if table_lines:
            table_text = "[表格]\n" + "\n".join(table_lines)
            paragraphs.append(table_text)

    numbered_text = _build_numbered_text(paragraphs)
    return {
        "text": numbered_text,
        "paragraphs": paragraphs,
        "headers": headers,
        "page_count": 0,  # DOCX doesn't have a native page count
    }


def _parse_markdown(content: bytes) -> dict:
    """Parse Markdown preserving heading hierarchy."""
    text = content.decode("utf-8", errors="replace")
    lines = text.split("\n")

    paragraphs = []
    headers = []
    current_para = []

    for line in lines:
        stripped = line.strip()

        # Detect headers
        header_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if header_match:
            # Flush current paragraph
            if current_para:
                paragraphs.append("\n".join(current_para))
                current_para = []

            level = len(header_match.group(1))
            header_text = header_match.group(2).strip()
            para_index = len(paragraphs)
            paragraphs.append(header_text)
            headers.append({"level": level, "text": header_text, "para_index": para_index})
        elif stripped == "":
            # Empty line = paragraph break
            if current_para:
                paragraphs.append("\n".join(current_para))
                current_para = []
        else:
            current_para.append(stripped)

    # Flush remaining
    if current_para:
        paragraphs.append("\n".join(current_para))

    numbered_text = _build_numbered_text(paragraphs)
    return {
        "text": numbered_text,
        "paragraphs": paragraphs,
        "headers": headers,
        "page_count": 0,
    }


def _build_numbered_text(paragraphs: list[str]) -> str:
    """Build text with paragraph numbers for LLM consumption."""
    lines = []
    for i, para in enumerate(paragraphs):
        lines.append(f"[{i}] {para}")
    return "\n\n".join(lines)


def extract_headers_summary(paragraphs: list[str], headers: list[dict]) -> str:
    """Extract first 2000 chars + all headers for long document structure sensing."""
    # Get first 2000 chars worth of paragraphs
    first_part = []
    char_count = 0
    for para in paragraphs:
        if char_count >= 2000:
            break
        first_part.append(para)
        char_count += len(para)

    # Collect all header lines
    header_lines = [f"{'#' * h['level']} {h['text']} (段落 {h['para_index']})" for h in headers]

    return (
        "=== 文档前 2000 字 ===\n"
        + "\n\n".join(f"[{i}] {p}" for i, p in enumerate(first_part))
        + "\n\n=== 全部标题 ===\n"
        + "\n".join(header_lines)
    )
